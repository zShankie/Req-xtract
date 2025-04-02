import os
import re
import PyPDF2
import spacy
import nltk
import pandas as pd
import streamlit as st
import torch
from transformers import DistilBertTokenizer, DistilBertForSequenceClassification
from nltk.tokenize import sent_tokenize
from datetime import datetime
import base64
from io import BytesIO
from docx import Document  

@st.cache_resource
def download_nltk_data():
    nltk.download('punkt', quiet=True)

@st.cache_resource
def load_spacy_model():
    try:
        return spacy.load("en_core_web_sm")
    except Exception:
        import subprocess
        subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
        return spacy.load("en_core_web_sm")

@st.cache_resource
def load_model():
    model_path = "fine_tuned_distilbert"  
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    tokenizer = DistilBertTokenizer.from_pretrained(model_path)
    model = DistilBertForSequenceClassification.from_pretrained(model_path).to(device)
    model.eval()
    return tokenizer, model, device



class RequirementsExtractor:
    def __init__(self):
        download_nltk_data()
        self.nlp = load_spacy_model()
        self.tokenizer, self.model, self.device = load_model()
        self.extracted_requirements = []  # List of dicts: { "Requirement Text": ..., "Predicted Type": ... }

    def extract_text_from_pdf(self, pdf_file):
        """Extract text from a PDF file with progress feedback."""
        reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        total_pages = len(reader.pages)
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for page_num in range(total_pages):
            status_text.text(f"Processing page {page_num+1}/{total_pages}...")
            page_text = reader.pages[page_num].extract_text() or ""
            text += page_text
            progress_bar.progress((page_num + 1) / total_pages)
            
        status_text.text(f"Successfully extracted {len(text)} characters from {total_pages} pages.")
        return text

    def predict_requirement(self, text):
        """Classify a requirement sentence as FR or NFR using the fine-tuned model."""
        encoding = self.tokenizer(text, truncation=True, padding="max_length", max_length=512, return_tensors="pt")
        input_ids = encoding["input_ids"].to(self.device)
        attention_mask = encoding["attention_mask"].to(self.device)
        
        with torch.no_grad():
            outputs = self.model(input_ids, attention_mask=attention_mask)
            logits = outputs.logits
            prediction = torch.argmax(logits, dim=1).item()
        # Return "FR" if prediction == 0 else "NFR" (adjust as needed based on your training)
        return "FR" if prediction == 0 else "NFR"

    def identify_requirements(self, text):
        """Identify and classify requirements from text."""
        st.info("Analyzing text for requirements...")
        sentences = sent_tokenize(text)
        st.text(f"Found {len(sentences)} sentences to analyze.")
        
        requirement_patterns = [
            r"(?i).*\b(shall|must|should|will|need to|has to|requires|required to)\b.*",
            r"(?i).*\b(necessary|mandatory|essential|important|critical)\b.*",
            r"(?i).*\b(system|application|platform|software|solution|user)\b.*\b(shall|must|should|will)\b.*",
            r"(?i).*\b(functionality|feature|capability)\b.*",
            r"(?i).*\b(ability to|enable|allow)\b.*",
        ]
        
        requirements = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, sentence in enumerate(sentences):
            if i % 10 == 0:
                status_text.text(f"Analyzed {i}/{len(sentences)} sentences...")
                progress_bar.progress(min((i + 1) / len(sentences), 1.0))
            clean_sentence = sentence.strip()
            if len(clean_sentence.split()) < 3:
                continue
            # Check against any requirement pattern
            is_requirement = any(re.match(pattern, clean_sentence) for pattern in requirement_patterns)
            if is_requirement:
                doc = self.nlp(clean_sentence)
                # Further filtering: check for key verbs
                if any(token.lemma_ in ["shall", "must", "should", "need", "require", "provide", "support", "allow", "enable"] for token in doc):

                    clean_sentence = re.sub(r"^(FR|NFR)-\d+\s*[:|-]?\s*", "", clean_sentence)
                    predicted_type = self.predict_requirement(clean_sentence)
                    requirements.append({
                        "Requirement Text": clean_sentence,
                        "Predicted Type": predicted_type
                    })
        
        if not requirements:
            st.warning("No requirements were extracted. Generating placeholder requirements.")
            requirements = [
                {"Requirement Text": "The system shall provide a user-friendly interface for managing tasks.", "Predicted Type": "FR"},
                {"Requirement Text": "The system shall ensure data security and compliance with industry standards.", "Predicted Type": "NFR"}
            ]
        
        status_text.text(f"Found {len(requirements)} potential requirements.")
        self.extracted_requirements = requirements
        return requirements

    def create_download_link_txt(self, requirements, filename="requirements.txt"):
        """Create a download link for the requirements as a text file."""
        content = "# EXTRACTED REQUIREMENTS\n"
        content += f"# Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        for i, req in enumerate(requirements, 1):
            content += f"REQ-{i:03d}: [{req['Predicted Type']}] {req['Requirement Text']}\n\n"
        b64 = base64.b64encode(content.encode()).decode()
        href = f'<a href="data:file/txt;base64,{b64}" download="{filename}">Download Requirements (.txt)</a>'
        return href

    def create_download_link_csv(self, requirements, filename="requirements.csv"):
        """Create a download link for the requirements as a CSV file."""
        df = pd.DataFrame({
            'ID': [f"REQ-{i:03d}" for i in range(1, len(requirements) + 1)],
            'Requirement Text': [req["Requirement Text"] for req in requirements],
            'Predicted Type': [req["Predicted Type"] for req in requirements]
        })
        csv = df.to_csv(index=False)
        b64 = base64.b64encode(csv.encode()).decode()
        href = f'<a href="data:file/csv;base64,{b64}" download="{filename}">Download Requirements (.csv)</a>'
        return href



def generate_srs_document_txt(fr_list, nfr_list, project_title="My Software Project"):
    """
    Generate an SRS document as a text string using provided FR and NFR lists.
    If none are provided, default placeholders are used.
    """
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    if not fr_list:
        fr_list = ["The system shall provide an intuitive user interface to manage operations.",
                   "The system shall allow real-time monitoring of system performance."]
    if not nfr_list:
        nfr_list = ["The system shall ensure a minimum uptime of 99.9%.",
                    "The system shall comply with modern security standards and encryption protocols."]
    
    srs = f"# Software Requirements Specification (SRS)\n\n"
    srs += f"**Project:** {project_title}\n"
    srs += f"**Date:** {now}\n\n"
    
    srs += "## 1. Introduction\n"
    srs += ("This document describes the Software Requirements Specification (SRS) for the project. "
            "It includes both Functional Requirements (FR) and Non‑Functional Requirements (NFR) which have been extracted and/or generated.\n\n")
    srs += "### 1.1 Purpose\n"
    srs += ("The purpose of this document is to provide a detailed description of the system to be developed and serve as a guideline for design, development, and testing teams.\n\n")
    srs += "### 1.2 Scope\n"
    srs += ("This SRS document outlines functionalities, performance criteria, design constraints, and quality attributes of the system. "
            "It covers all major aspects of the software product.\n\n")
    
    srs += "## 2. Overall Description\n"
    srs += ("This section provides an overview of the system including user characteristics, assumptions, and constraints that affect system design.\n\n")
    
    srs += "## 3. Functional Requirements (FR)\n"
    srs += "The following table lists the Functional Requirements (FR):\n\n"
    fr_df = pd.DataFrame({
        "ID": [f"FR-{i:03d}" for i in range(1, len(fr_list) + 1)],
        "Requirement": fr_list
    })
    srs += fr_df.to_markdown(index=False) + "\n\n"
    
    srs += "## 4. Non-Functional Requirements (NFR)\n"
    srs += "The following table lists the Non‑Functional Requirements (NFR):\n\n"
    nfr_df = pd.DataFrame({
        "ID": [f"NFR-{i:03d}" for i in range(1, len(nfr_list) + 1)],
        "Requirement": nfr_list
    })
    srs += nfr_df.to_markdown(index=False) + "\n\n"
    
    srs += "## 5. External Interface Requirements\n"
    srs += ("This section describes the external interfaces of the system.\n\n"
            "### 5.1 User Interfaces\n"
            "The system will provide a web-based user interface supporting modern browsers (Chrome, Firefox, Edge, Safari) "
            "and a mobile interface optimized for both Android and iOS devices. It will use responsive design to ensure usability across devices.\n\n"
            "### 5.2 Hardware Interfaces\n"
            "The system is designed to run on standard server hardware and may interface with peripheral devices such as printers, barcode scanners, or IoT devices if needed.\n\n"
            "### 5.3 Software Interfaces\n"
            "The system will integrate with external systems via RESTful APIs. It will support JSON and XML data formats and offer secure endpoints for third-party integrations.\n\n")
    
    srs += "## 6. Appendix & References\n"
    srs += ("### 6.1 Glossary\n"
            "Key terms and acronyms used in this document include:\n"
            "- UI: User Interface\n"
            "- API: Application Programming Interface\n"
            "- DB: Database\n\n"
            "### 6.2 Supporting Documentation\n"
            "Supporting documentation may include software design documents, user manuals, and industry best practices.\n\n"
            "### 6.3 References\n"
            "1. IEEE Standard for Software Requirements Specifications (IEEE 830)\n"
            "2. ISO/IEC/IEEE 29148:2018 Systems and software engineering — Requirements engineering\n"
            "3. Relevant industry publications and white papers.\n\n")
    
    return srs

def generate_srs_document_docx(fr_list, nfr_list, project_title="My Software Project"):
    """
    Generate an SRS document as a DOCX file using provided FR and NFR lists.
    """
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    if not fr_list:
        fr_list = ["The system shall provide an intuitive user interface to manage operations.",
                   "The system shall allow real-time monitoring of system performance."]
    if not nfr_list:
        nfr_list = ["The system shall ensure a minimum uptime of 99.9%.",
                    "The system shall comply with modern security standards and encryption protocols."]
    
    document = Document()
    document.add_heading("Software Requirements Specification (SRS)", 0)
    document.add_paragraph(f"Project: {project_title}")
    document.add_paragraph(f"Date: {now}")
    
    document.add_heading("1. Introduction", level=1)
    document.add_paragraph("This document describes the Software Requirements Specification (SRS) for the project. It includes both Functional Requirements (FR) and Non‑Functional Requirements (NFR) which have been extracted and/or generated.")
    
    document.add_heading("1.1 Purpose", level=2)
    document.add_paragraph("The purpose of this document is to provide a detailed description of the system to be developed and serve as a guideline for design, development, and testing teams.")
    
    document.add_heading("1.2 Scope", level=2)
    document.add_paragraph("This SRS document outlines functionalities, performance criteria, design constraints, and quality attributes of the system. It covers all major aspects of the software product.")
    
    document.add_heading("2. Overall Description", level=1)
    document.add_paragraph("This section provides an overview of the system including user characteristics, assumptions, and constraints that affect system design.")
    
    document.add_heading("3. Functional Requirements (FR)", level=1)
    table_fr = document.add_table(rows=1, cols=2)
    hdr_cells = table_fr.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "Requirement"
    for i, req in enumerate(fr_list, 1):
        row_cells = table_fr.add_row().cells
        row_cells[0].text = f"FR-{i:03d}"
        row_cells[1].text = req
    
    document.add_heading("4. Non-Functional Requirements (NFR)", level=1)
    table_nfr = document.add_table(rows=1, cols=2)
    hdr_cells = table_nfr.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "Requirement"
    for i, req in enumerate(nfr_list, 1):
        row_cells = table_nfr.add_row().cells
        row_cells[0].text = f"NFR-{i:03d}"
        row_cells[1].text = req
    
    document.add_heading("5. External Interface Requirements", level=1)
    document.add_heading("5.1 User Interfaces", level=2)
    document.add_paragraph("The system will provide a web-based user interface supporting modern browsers (Chrome, Firefox, Edge, Safari) and a mobile interface optimized for both Android and iOS devices. The design will be responsive to various screen sizes.")
    document.add_heading("5.2 Hardware Interfaces", level=2)
    document.add_paragraph("The system is designed to operate on standard server hardware and may interface with peripheral devices such as printers, barcode scanners, or IoT devices when required.")
    document.add_heading("5.3 Software Interfaces", level=2)
    document.add_paragraph("The system will integrate with external systems via RESTful APIs supporting JSON and XML data formats. It will also offer secure endpoints for third-party integrations.")
    
    document.add_heading("6. Appendix & References", level=1)
    document.add_heading("6.1 Glossary", level=2)
    document.add_paragraph("Key terms and acronyms:\n- UI: User Interface\n- API: Application Programming Interface\n- DB: Database")
    document.add_heading("6.2 Supporting Documentation", level=2)
    document.add_paragraph("Supporting documentation includes software design documents, user manuals, industry standards, and reference architectures.")
    document.add_heading("6.3 References", level=2)
    document.add_paragraph("1. IEEE Standard for Software Requirements Specifications (IEEE 830)\n"
                           "2. ISO/IEC/IEEE 29148:2018 Requirements Engineering\n"
                           "3. Additional industry publications and white papers.")
    
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def create_download_link_srs(text_content, filename="SRS_document.txt"):
    """Create a download link for the SRS document as a TXT file."""
    b64 = base64.b64encode(text_content.encode()).decode()
    href = f'<a href="data:file/txt;base64,{b64}" download="{filename}">Download SRS Document (.txt)</a>'
    return href

def create_download_link_srs_docx(docx_buffer, filename="SRS_document.docx"):
    """Create a download link for the SRS document as a DOCX file."""
    b64 = base64.b64encode(docx_buffer.read()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download SRS Document (.docx)</a>'
    return href

# Main Streamlit App

def main():
    st.set_page_config(
        page_title="Requirements Extraction, Classification & SRS Generator",
        page_icon="📝",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS for styling
    st.markdown("""
    <style>
        .main { background-color: #f8f9fa; }
        .stButton>button {
            background-color: #4CAF50; color: white; border-radius: 5px;
            padding: 0.5rem 1rem; border: none; font-weight: bold;
        }
        .stButton>button:hover { background-color: #45a049; }
        .stDownloadButton>button {
            background-color: #2196F3; color: white; border-radius: 5px;
            padding: 0.5rem 1rem; border: none; font-weight: bold;
        }
        .stDownloadButton>button:hover { background-color: #0b7dda; }
        .stTextArea>div>div>textarea {
            background-color: #000000; color: #ffffff;
            border: 1px solid #ced4da; border-radius: 5px;
        }
        .stDataFrame { border-radius: 5px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
        .sidebar .sidebar-content { background-color: #343a40; color: white; }
        h1, h2, h3 { color: #2c3e50; }
        .success-box { background-color: #d4edda; color: #155724; padding: 1rem; border-radius: 5px; margin-bottom: 1rem; }
        .info-box { background-color: #d1ecf1; color: #0c5460; padding: 1rem; border-radius: 5px; margin-bottom: 1rem; }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div style="background-color: #2c3e50; padding: 2rem; border-radius: 10px; margin-bottom: 2rem;">
        <h1 style="color: white; margin: 0;">📝 Requirements Extraction, Classification & SRS Generator</h1>
        <p style="color: #bdc3c7; margin: 0.5rem 0 0;">Extract requirements from PDFs, classify as FR/NFR, and generate an SRS document</p>
    </div>
    """, unsafe_allow_html=True)
    

    if 'pdf_text' not in st.session_state:
        st.session_state.pdf_text = None
    if 'requirements' not in st.session_state:
        st.session_state.requirements = []
    if 'file_name' not in st.session_state:
        st.session_state.file_name = None
    
    extractor = RequirementsExtractor()
    
    # Sidebar for PDF upload and extraction
    with st.sidebar:
        st.markdown("""
        <div style="padding: 1rem; background-color: #2c3e50; border-radius: 10px; margin-bottom: 2rem;">
            <h3 style="color: white;">Upload Document</h3>
        </div>
        """, unsafe_allow_html=True)
        
        uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"], label_visibility="collapsed")
        
        if uploaded_file is not None:
            st.markdown("""
            <div style="padding: 1rem; background-color: #2c3e50; border-radius: 10px; margin: 1rem 0;">
                <h3 style="color: white;">Document Info</h3>
            </div>
            """, unsafe_allow_html=True)
            st.write(f"**File name:** {uploaded_file.name}")
            st.write(f"**File size:** {uploaded_file.size / 1024:.2f} KB")
            
            if st.button("Extract & Classify Requirements", use_container_width=True):
                with st.spinner("Processing document..."):
                    st.session_state.file_name = uploaded_file.name
                    st.session_state.pdf_text = extractor.extract_text_from_pdf(uploaded_file)
                    st.session_state.requirements = extractor.identify_requirements(st.session_state.pdf_text)
        
        st.markdown("---")
        st.markdown("""
        <div style="padding: 1rem; background-color: #2c3e50; border-radius: 10px; margin: 1rem 0;">
            <h3 style="color: white;">About</h3>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("""
        This tool extracts requirements from PDF documents, classifies them as Functional (FR) or Non‑Functional (NFR), 
        and generates a Software Requirements Specification (SRS) document with generic, hardcoded content where needed.
        """)
        st.markdown("Created with ❤️ by Team Hawk")
    
    # Main content columns: Text preview and Requirements table
    col1, col2 = st.columns(2, gap="large")
    
    with col1:
        st.markdown("""
        <div style="padding: 1rem; background-color: #2c3e50; border-radius: 10px; margin-bottom: 1rem;">
            <h2 style="color: white;">Document Text</h2>
        </div>
        """, unsafe_allow_html=True)
        if st.session_state.pdf_text:
            preview_text = st.session_state.pdf_text[:5000] + ("..." if len(st.session_state.pdf_text) > 5000 else "")
            st.text_area("Extracted Text Preview", preview_text, height=400, label_visibility="collapsed")
            st.caption(f"📄 Showing first 5,000 characters of {len(st.session_state.pdf_text):,} total")
        else:
            st.markdown("""
            <div class="info-box">
                <h3>No document loaded</h3>
                <p>Upload a PDF using the sidebar and click "Extract & Classify Requirements" to begin.</p>
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div style="padding: 1rem; background-color: #2c3e50; border-radius: 10px; margin-bottom: 1rem;">
            <h2 style="color: white;">Extracted Requirements</h2>
        </div>
        """, unsafe_allow_html=True)
        if st.session_state.requirements:
            st.markdown(f"""
            <div class="success-box">
                <h3>Success!</h3>
                <p>Found {len(st.session_state.requirements)} potential requirements.</p>
            </div>
            """, unsafe_allow_html=True)
            
            df = pd.DataFrame({
                'ID': [f"REQ-{i:03d}" for i in range(1, len(st.session_state.requirements) + 1)],
                'Requirement': [req["Requirement Text"] for req in st.session_state.requirements],
                'Predicted Type': [req["Predicted Type"] for req in st.session_state.requirements]
            })
            st.dataframe(df, height=400, use_container_width=True)
            
            st.markdown("---")
            st.markdown("### Download Requirements")
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            base_filename = st.session_state.file_name.replace('.pdf', '') if st.session_state.file_name else "requirements"
            txt_filename = f"{base_filename}_requirements_{timestamp}.txt"
            csv_filename = f"{base_filename}_requirements_{timestamp}.csv"
            
            dl_col1, dl_col2 = st.columns(2)
            with dl_col1:
                st.markdown(extractor.create_download_link_txt(st.session_state.requirements, txt_filename), unsafe_allow_html=True)
            with dl_col2:
                st.markdown(extractor.create_download_link_csv(st.session_state.requirements, csv_filename), unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="info-box">
                <h3>No requirements extracted yet</h3>
                <p>After processing a document, extracted requirements (with FR/NFR classification) will appear here.</p>
            </div>
            """, unsafe_allow_html=True)
    
    # SRS Generation Section
    st.markdown("---")
    st.markdown("""
    <div style="padding: 1rem; background-color: #2c3e50; border-radius: 10px; margin-bottom: 1rem;">
        <h2 style="color: white;">Generated SRS Document</h2>
    </div>
    """, unsafe_allow_html=True)
    
    if st.session_state.requirements:
        fr_list = [req["Requirement Text"] for req in st.session_state.requirements if req["Predicted Type"] == "FR"]
        nfr_list = [req["Requirement Text"] for req in st.session_state.requirements if req["Predicted Type"] == "NFR"]
        
        srs_txt = generate_srs_document_txt(fr_list, nfr_list, project_title="Automated Workforce Management")
        st.text_area("SRS Document (TXT)", srs_txt, height=500, label_visibility="collapsed")
        st.markdown(create_download_link_srs(srs_txt), unsafe_allow_html=True)
        
        # Generate DOCX version and provide download link
        docx_buffer = generate_srs_document_docx(fr_list, nfr_list, project_title="Example Software Project")
        st.markdown(create_download_link_srs_docx(docx_buffer), unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="info-box">
            <h3>SRS Document Not Generated</h3>
            <p>Please extract requirements from a PDF to generate an SRS document.</p>
        </div>
        """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()